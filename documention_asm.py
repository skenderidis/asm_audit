import os
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import RGBColor
import datetime
import requests
from requests.auth import HTTPDigestAuth
import json
import getpass

requests.packages.urllib3.disable_warnings() 

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width
     

def word_file_results(document, results, overview, suggestions):
	num_of_suggestions = 0
	document.add_heading(overview['name'], level=2)
	document.add_paragraph('The following section analyzing the configuration of the ASM policy')
	document.add_heading('Overview', level=3)

	document.add_paragraph('The following table provides a quick view on the entities configured this policy and their enforced status')
	
	table = document.add_table(rows=8, cols=3)
	table.style = 'Table Grid'
	table.cell(0,0).text = 'Entities'
	table.cell(1,0).text = 'File Types'
	table.cell(2,0).text = 'URLs'
	table.cell(3,0).text = 'Parameters'
	table.cell(4,0).text = 'Signatures'
	table.cell(5,0).text = 'Cookies'
	table.cell(6,0).text = 'HTTP Compliance'
	table.cell(7,0).text = 'Evasion'
	#------------------------------------------------------------------
	table.cell(0,1).text = 'Total Configured'
	table.cell(1,1).text = str(results['file_type_total'])
	table.cell(2,1).text = str(results['urls_total'])
	table.cell(3,1).text = str(results['param_total'])
	table.cell(4,1).text = str(results['sig_total'])
	table.cell(5,1).text = str(results['cookies_total'])
	table.cell(6,1).text = str(results['compliance_total'])
	table.cell(7,1).text = str(results['evasion_total'])
	#------------------------------------------------------------------
	table.cell(0,2).text = 'Not Enforced'
	table.cell(1,2).text = str(results['file_type_not_enforced'])
	table.cell(2,2).text = str(results['urls_not_enforced'])
	table.cell(3,2).text = str(results['param_not_enforced'])
	table.cell(4,2).text = str(results['sig_not_enforced'])
	table.cell(5,2).text = str(results['cookies_not_enforced'])
	table.cell(6,2).text = str(results['compliance_not_enforced'])
	table.cell(7,2).text = str(results['evasion_not_enforced'])


	document.add_paragraph()
	document.add_paragraph('The following table provides the summary of all the recommendations of the health check for this ASM policy.')
	
	num_of_suggestions = 0
	table = document.add_table(rows=1, cols=4)

	table.style = 'Table Grid'
	table.cell(0,0).text = '#'
	table.cell(0,1).text = 'Severity'
	table.cell(0,2).text = 'Suggestion'
	table.cell(0,3).text = 'Category'
	
	for key in suggestions:
		if key['severity'] == "error":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']

	for key in suggestions:
		if key['severity'] == "warning":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']
	
	for key in suggestions:
		if key['severity'] == "info":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']
				
	set_column_width(table.columns[0], Cm(1))		
	set_column_width(table.columns[1], Cm(1))		
	set_column_width(table.columns[2], Cm(12))		
	set_column_width(table.columns[3], Cm(3))		

	document.add_paragraph()
	document.add_paragraph()	
	
	
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_overview (document, overview, suggestions):
	violations = ["IP is blacklisted","Malformed XML data","Malformed JSON data","Disallowed file upload content detected","Virus detected","Brute Force: Maximum login attempts are exceeded"]
	num_of_suggestions = 0
	document.add_heading('General Settings', level=3)
	
	table = document.add_table(rows=17, cols=2)
	table.style = 'Table Grid'
	table.cell(0,0).text = 'Settings'
	table.cell(1,0).text = 'Policy Name'
	table.cell(2,0).text = 'Partition'
	table.cell(3,0).text = 'Enforcement mode'
	table.cell(4,0).text = 'Applied to vServers'
	table.cell(5,0).text = 'Application Language'
	table.cell(6,0).text = 'Brute force Protection'
	table.cell(7,0).text = 'DataGuard'
	table.cell(8,0).text = 'Antivirus'
	table.cell(9,0).text = 'Created By'
	table.cell(10,0).text = 'Created Date'
	table.cell(11,0).text = 'Last Updated'
	table.cell(12,0).text = 'Policy is case sensitive'
	table.cell(13,0).text = 'Mask Credit Card Numbers in Request Log'
	table.cell(14,0).text = 'Trust XFF'
	table.cell(15,0).text = 'Custom XFF'		
	table.cell(16,0).text = 'Trigger ASM iRule Events'		
	#------------------------------------------------------------------
	table.cell(0,1).text = 'Values'
	table.cell(1,1).text = overview['name']
	table.cell(2,1).text = overview['partition']
	table.cell(3,1).text = overview['enforcementMode']
	table.cell(4,1).text = '\n'.join(overview['virtualServers'])
	table.cell(5,1).text = overview['applicationLanguage']
	table.cell(6,1).text = overview['brute_enabled'] + " (on " + str(overview['Login_pages_totalItems']) + " login pages)"
	table.cell(7,1).text = overview['data_guard_enabled']
	table.cell(8,1).text = overview['inspectHttpUploads']
	table.cell(9,1).text = overview['creatorName']
	table.cell(10,1).text = overview['createdDatetime']		
	table.cell(11,1).text = overview['lastUpdateMicros']		
	table.cell(12,1).text = overview['caseInsensitive']		
	table.cell(13,1).text = overview['maskCreditCardNumbersInRequest']		
	table.cell(14,1).text = overview['trustXff']		
	table.cell(15,1).text = '\n'.join(overview['customXffHeaders'])		
	table.cell(16,1).text = overview['triggerAsmIruleEvent']		

	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Overview":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Overview":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Overview":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()
	document.save("reports/F5 ASM - Config Review.docx")
	
def word_file_learning (document, policy_builder, whitelist, suggestions):
	num_of_suggestions = 0
	document.add_heading('Learning Configuration', level=3)
	document.add_paragraph('The following table shows the Learning configuration of the ASM Policy')
	table = document.add_table(rows=21, cols=2)
	table.style = 'Table Grid'
	table.cell(0,0).text = 'Settings'
	table.cell(1,0).text = 'Learning Mode'	
	table.cell(2,0).text = 'Trust All IPs'
	table.cell(3,0).text = 'Trusted sources for learning'
	table.cell(4,0).text = 'Trusted hours for learning'
	table.cell(5,0).text = 'Untrusted sources for learning'
	table.cell(6,0).text = 'Untrusted hours for learning'
	table.cell(7,0).text = 'Full Inspection'
	table.cell(8,0).text = 'Learn Inactive Entities'
	table.cell(9,0).text = 'Learn New File Types'
	table.cell(10,0).text = 'Max Learned File Types'
	table.cell(11,0).text = 'Learn New URLs'
	table.cell(12,0).text = 'Max Learned URLs'
	table.cell(13,0).text = 'Learn New Parameters'
	table.cell(14,0).text = 'Max Learned Parameters'
	table.cell(15,0).text = 'Parameter Learning Level'
	table.cell(16,0).text = 'Learn Integer Values'
	table.cell(17,0).text = 'Classify Value Content'
	table.cell(18,0).text = 'Learn New Cookies'
	table.cell(19,0).text = 'Max Learned Cookies'
	table.cell(20,0).text = 'Learn Redirection Domains'
	#------------------------------------------------------------------
	table.cell(0,1).text = 'Values'
	table.cell(1,1).text = policy_builder['learningMode']
	table.cell(2,1).text = policy_builder['trustAllIps']
	table.cell(3,1).text = str(policy_builder['trusted_loosen_source'])
	table.cell(4,1).text = str(policy_builder['trusted_loosen_hours'])
	table.cell(5,1).text = str(policy_builder['untrusted_loosen_source'])
	table.cell(6,1).text = str(policy_builder['untrusted_loosen_hours'])
	table.cell(7,1).text = policy_builder['enableFullPolicyInspection']		
	table.cell(8,1).text = policy_builder['learnInactiveEntities']		
	table.cell(9,1).text = policy_builder['learnExplicitFiletypes']		
	table.cell(10,1).text = str(policy_builder['maximumFileTypes'])
	table.cell(11,1).text = policy_builder['learnExplicitUrls']		
	table.cell(12,1).text = str(policy_builder['maximumUrls'])		
	table.cell(13,1).text = policy_builder['learnExplicitParameters']		
	table.cell(14,1).text = str(policy_builder['maximumParameters'])		
	table.cell(15,1).text = policy_builder['parameterLearningLevel']		
	table.cell(16,1).text = policy_builder['parametersIntegerValue']		
	table.cell(17,1).text = policy_builder['classifyParameters']		
	table.cell(18,1).text = policy_builder['learnExplicitCookies']		
	table.cell(19,1).text = str(policy_builder['maximumCookies'])	
	table.cell(20,1).text = policy_builder['learnExplicitRedirectionDomains']		

	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Policy Builder":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Policy Builder":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Policy Builder":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()
			
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_compliance (document, blocking_settings, compliance, suggestions):

	violations = ["HTTP protocol compliance failed"]
	num_of_suggestions = 0
	document.add_heading('Compliance Violations', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for HTTP Compliance protection.')
	
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']

	set_column_width(table.columns[0], Cm(9.5))			

	document.add_paragraph()
	document.add_paragraph('The following section summarizes the configuration for HTTP Compliance protection.')

	table = document.add_table(rows=1, cols=3)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'HTTP Compliance Violation'
	table.cell(0,1).text = 'Enabled'
	table.cell(0,2).text = 'Learn'
	

	for key in compliance:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['enabled']
		cells[2].text = key['learn']

	set_column_width(table.columns[0], Cm(12.5))			

	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="HTTP Compliance":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="HTTP Compliance":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="HTTP Compliance":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')


	document.add_paragraph()
	document.add_paragraph()
	
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_evasion (document, blocking_settings, evasions, suggestions):
	
	violations = ["Evasion technique detected"]
	num_of_suggestions = 0
	document.add_heading('Evasion Techniques', level=3)
	document.add_paragraph('The Evasion technique detected violation is triggered when the BIG-IP ASM system fails to normalize requests. Normalization is the process of decoding requests that are encoded. The system needs to perform normalization because some applications send requests that contain different types of encoded escapes that the BIG-IP ASM system needs to interpret before doing any further processing')

	document.add_paragraph('The following section shows the blocking settings configuration for Evasion protection.')
	
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']

	set_column_width(table.columns[0], Cm(9.5))			
				
	document.add_paragraph()
	document.add_paragraph('The following section summarizes the configuration for Evasion techiques protection.')

	table = document.add_table(rows=1, cols=3)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Evasion Techniques'
	table.cell(0,1).text = 'Enabled'
	table.cell(0,2).text = 'Learn'
	for key in evasions:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['enabled']
		cells[2].text = key['learn']

	set_column_width(table.columns[0], Cm(12.5))			
		

	document.add_paragraph()
	document.add_paragraph()
				
	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Evasion":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Evasion":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Evasion":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')

	
	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 ASM - Config Review.docx")

def word_file_signatures (document, signatures_overview, signature_sets, urls, headers, parameters, cookies, suggestions):
	num_of_suggestions = 0
	document.add_heading('Signatures', level=3)
	document.add_paragraph('This section of the report summarizes the configuration done for Attack Signature protection.')

	
	p = document.add_paragraph()
	runner = p.add_run("Signature Staging:		")
	runner.bold = True
	p.add_run (signatures_overview['signatureStaging'])
	p = document.add_paragraph()
	runner = p.add_run("Place New Signature in Staging:	")
	runner.bold = True
	p.add_run (signatures_overview['placeSignaturesInStaging'])
	document.add_paragraph()
	
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Signature Set Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'

	for key in signature_sets:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['learn']
		cells[2].text = key['alarm']
		cells[3].text = key['block']
				
	document.add_paragraph()
	p = document.add_paragraph()
	runner = p.add_run("Total Signatures:			")
	runner.bold = True
	p.add_run (str(signatures_overview['total']))
	p = document.add_paragraph()
	runner = p.add_run("Staging Signatures:		")
	runner.bold = True
	p.add_run (str(signatures_overview['staging']))
	p = document.add_paragraph()
	runner = p.add_run("Disabled Signatures:		")
	runner.bold = True
	p.add_run (str(signatures_overview['enabled']))				
	p = document.add_paragraph()
	runner = p.add_run("Latest Signature update:		")
	runner.bold = True
	p.add_run (signatures_overview['latest_sig_update'])	
	document.add_paragraph()

		
	document.add_heading('Signature Overrides', level=4)
	document.add_paragraph('This section shows all entities that have Signatures Overrides and therefore these attack signatures are not being applied.')

	
	table = document.add_table(rows=1, cols=3)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Type'
	table.cell(0,1).text = 'Name'
	table.cell(0,2).text = 'Signature Overrides'

	
	for key in urls:
		if len(key['signatureOverrides'])>1  and key['attackSignaturesCheck'] == "Yes":
			cells = table.add_row().cells
			cells[0].text = "URL"
			cells[1].text = key['name']
			cells[2].text = '\n'.join(key['signatureOverrides'])
		else:
			if (key['signatureOverrides'][0]!="None")  and key['attackSignaturesCheck'] == "Yes":
				cells = table.add_row().cells
				cells[0].text = "URL"
				cells[1].text = key['name']
				cells[2].text = '\n'.join(key['signatureOverrides'])

	for key in parameters:
		if len(key['signatureOverrides'])>1 and key['attackSignaturesCheck'] == "Yes":
			cells = table.add_row().cells
			cells[0].text = "Parameter"
			cells[1].text = key['name']
			cells[2].text = '\n'.join(key['signatureOverrides'])
		else:
			if (key['signatureOverrides'][0]!="None") and key['attackSignaturesCheck'] == "Yes":
				cells = table.add_row().cells
				cells[0].text = "Parameter"
				cells[1].text = key['name']
				cells[2].text = '\n'.join(key['signatureOverrides'])

	for key in headers:
		if len(key['signatureOverrides'])>1 and key['checkSignatures'] == "Yes":
			cells = table.add_row().cells
			cells[0].text = "Header"
			cells[1].text = key['name']
			cells[2].text = '\n'.join(key['signatureOverrides'])
		else:
			if (key['signatureOverrides'][0]!="None") and key['checkSignatures'] == "Yes":
				cells = table.add_row().cells
				cells[0].text = "Header"
				cells[1].text = key['name']
				cells[2].text = '\n'.join(key['signatureOverrides'])

	for key in cookies:
		if len(key['signatureOverrides'])>1 and key['attackSignaturesCheck'] == "Yes":
			cells = table.add_row().cells
			cells[0].text = "Cookie"
			cells[1].text = key['name']
			cells[2].text = '\n'.join(key['signatureOverrides'])
		else:
			if (key['signatureOverrides'][0]!="None") and key['attackSignaturesCheck'] == "Yes":
				cells = table.add_row().cells
				cells[0].text = "Cookie"
				cells[1].text = key['name']
				cells[2].text = '\n'.join(key['signatureOverrides'])

	set_column_width(table.columns[0], Cm(2))			
	set_column_width(table.columns[2], Cm(7.5))	
	set_column_width(table.columns[2], Cm(9))			
		
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Signatures":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Signatures":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Signatures":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')

	document.add_paragraph()
	document.add_paragraph()
	
	
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_types (document, file_types, blocking_settings, policy_builder, suggestions):			
	num_of_suggestions = 0
	violations = ["Illegal query string length", "Illegal request length", "Illegal POST data length", "Illegal URL length", "Illegal file type"]
	document.add_heading('File Types', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for File Type and File Type Length violations.')

	p = document.add_paragraph()
	runner = p.add_run("Learn New File Types:	")
	runner.bold = True
	p.add_run (policy_builder['learnExplicitFiletypes'])
	p = document.add_paragraph()
	runner = p.add_run("Max Learned File Types:	")
	runner.bold = True
	p.add_run (str(policy_builder['maximumFileTypes']))
	
	if (policy_builder['learnExplicitUrls'] != "always"):
			suggestions.append({'severity':'low', 'txt':' We recommend that you configure "Learn New File Types" to "Always"'})	

			
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']
				
	set_column_width(table.columns[0], Cm(9.5))			
	document.add_paragraph()		

	document.add_paragraph('The following section summarizes the configuration for File Type and File Type Length violations.')	
	table = document.add_table(rows=1, cols=6)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'File Types Extension'
	table.cell(0,1).text = 'Staging'
	table.cell(0,2).text = 'URL Length'
	table.cell(0,3).text = 'Query Length'
	table.cell(0,4).text = 'POST Length'
	table.cell(0,5).text = 'Request Length'

	for key in file_types:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['performStaging']
		cells[2].text = str(key['urlLength'])
		cells[3].text = str(key['queryStringLength'])
		cells[4].text = str(key['postDataLength'])
		cells[5].text = str(key['requestLength'])
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="File Types":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="File Types":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="File Types":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
				
	document.add_paragraph()
	document.add_paragraph()


	document.save("reports/F5 ASM - Config Review.docx")

def word_file_urls (document, overview, urls, blocking_settings, policy_builder, suggestions):

	violations = ["---- Nothing enabled ----"]
	num_of_suggestions = 0
	document.add_heading('URLs', level=3)
	document.add_paragraph('This section of the report summarizes the configuration done for URL protection.')

	p = document.add_paragraph()
	runner = p.add_run("Learn New URLs:		")
	runner.bold = True
	p.add_run (policy_builder['learnExplicitUrls'])
	p = document.add_paragraph()
	runner = p.add_run("Max Learned URLs:	")
	runner.bold = True
	p.add_run (str(policy_builder['maximumUrls']))
	p = document.add_paragraph()
	
	if (len(urls)<50):
		table = document.add_table(rows=1, cols=7)
		table.style = 'Table_rows'
		table.cell(0,0).text = 'URL'
		table.cell(0,1).text = 'Protocol'
		table.cell(0,2).text = 'Staging'
		table.cell(0,3).text = 'Check Signatures'
		table.cell(0,4).text = 'Check Meta-Char'
		table.cell(0,5).text = 'Signature Overides'
		table.cell(0,6).text = 'Meta-Char Overides'

		for key in urls:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['protocol']
			cells[2].text = key['performStaging']
			cells[3].text = key['attackSignaturesCheck']
			cells[4].text = key['metacharsOnUrlCheck']
			cells[5].text = '\n'.join(key['signatureOverrides'])
			cells[6].text = key['metacharOverrides']

	else:
		document.add_paragraph('There are '+ str(len(urls)) + ' URLs configured for this policy. Due to this high number we are not providing a detailed table and if you need to review individually each URL, please review them on the Excel Sheet provided.')
	
	
		
	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="URLs":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="URLs":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="URLs":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')

		
	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 ASM - Config Review.docx")

def word_file_parameters (document, parameters, blocking_settings, policy_builder, sensitive_param, suggestions):	
	
	num_of_suggestions = 0
	document.add_heading('Parameters', level=3)
	violations = ["Disallowed file upload content detected"]
	staging = 0
	disabled = 0
	document.add_paragraph('This section of the report summarizes the configuration done for Parameter protection.')

	p = document.add_paragraph()
	runner = p.add_run("Learn New Parameters:				")
	runner.bold = True
	p.add_run (policy_builder['learnExplicitParameters'])
	p = document.add_paragraph()
	runner = p.add_run("Max Learned Parameters:				")
	runner.bold = True
	p.add_run (str(policy_builder['maximumParameters']))
	p = document.add_paragraph()
	runner = p.add_run("Parameter Learning Level:			")
	runner.bold = True
	p.add_run (policy_builder['parameterLearningLevel'])
	p = document.add_paragraph()
	runner = p.add_run("Learn Integer Parameters:			")
	runner.bold = True
	p.add_run (policy_builder['parametersIntegerValue'])
	p = document.add_paragraph()
	runner = p.add_run("Classify Value Content of Learned Parameters:	")
	runner.bold = True
	p.add_run (policy_builder['classifyParameters'])
	document.add_paragraph()		


	if (len(parameters)<50):
		table = document.add_table(rows=1, cols=7)
		table.style = 'Table_rows'
		table.cell(0,0).text = 'Parameter'
		table.cell(0,1).text = 'Enforcement'
		table.cell(0,2).text = 'Staging'
		table.cell(0,3).text = 'Check Signatures'
		table.cell(0,4).text = 'Check Meta-Char'
		table.cell(0,5).text = 'Signature Overides'
		table.cell(0,6).text = 'Meta-Char Overides'

		for key in parameters:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['enforcement']
			cells[2].text = key['performStaging']
			cells[3].text = key['attackSignaturesCheck']
			cells[4].text = key['metacharsOnParameterValueCheck']
			cells[5].text = '\n'.join(key['signatureOverrides'])
			cells[6].text = key['valueMetacharOverrides']
			
			if (key['attackSignaturesCheck'] == "No"):
					disabled += 1
			if (key['performStaging'] == "Yes"):
					staging +=1	
	else:
		document.add_paragraph('There are '+ str(len(parameters)) + ' Parameters configured for this policy. Due to this high number we are not providing a detailed table and if you need to review individually each parameter, please review them on the Excel Sheet provided.')

	document.add_paragraph()
	document.add_paragraph('The following section summarizes the configuration for Sensitive Parameters.')
	table = document.add_table(rows=1, cols=1)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Sensitive Parameters'
	for key in sensitive_param:
		cells = table.add_row().cells
		cells[0].text = key['name']

		
	document.add_paragraph()
	document.add_paragraph()
	
	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Parameters":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Parameters":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Parameters":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')

			
	document.add_paragraph()
	document.add_paragraph()
	

	document.save("reports/F5 ASM - Config Review.docx")

def word_file_headers (document, overview, headers, blocking_settings, suggestions):

	violations = ["Illegal header length"]
	num_of_suggestions = 0
		
	document.add_heading('Headers', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for HTTP Header protection.')
	p = document.add_paragraph()
	runner = p.add_run("Max HTTP Header Length:	")
	runner.bold = True
	p.add_run (overview['maximumHttpHeaderLength'])

		
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']
				
	set_column_width(table.columns[0], Cm(9.5))		
	document.add_paragraph()
	

	document.add_paragraph('The following section summarizes the configuration done for Header protection.')
	
	table = document.add_table(rows=1, cols=7)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Check Signatures'
	table.cell(0,2).text = 'Signature Overrides'
	table.cell(0,3).text = 'Evasion Techniques'
	table.cell(0,4).text = 'URL Normalization'
	table.cell(0,5).text = 'Percent Decoding'
	table.cell(0,6).text = 'HTML Normalization'

	
	for key in headers:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['checkSignatures']
		cells[2].text = '\n'.join(key['signatureOverrides'])
		cells[3].text = key['normalizationViolations']
		cells[4].text = key['urlNormalization']
		cells[5].text = key['percentDecoding']
		cells[6].text = key['htmlNormalization']
			
	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Headers":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Headers":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Headers":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')

	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 ASM - Config Review.docx")

def word_file_cookies (document, overview, cookies, blocking_settings, policy_builder, suggestions):	

	violations = ["Illegal cookie length", "Modified ASM cookie", "Modified domain cookie(s)", "Cookie not RFC-compliant"]
	num_of_suggestions = 0

	document.add_heading('Cookies', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for Cookies protection.')

	p = document.add_paragraph()
	runner = p.add_run("Max Cookie Length:	")
	runner.bold = True
	p.add_run (overview['maximumCookieHeaderLength'])
	p = document.add_paragraph()
	runner = p.add_run("Learn New Cookies:	")
	runner.bold = True
	p.add_run (policy_builder['learnExplicitCookies'])
	p = document.add_paragraph()
	runner = p.add_run("Max Learned Cookies:	")
	runner.bold = True
	p.add_run (str(policy_builder['maximumCookies']))


	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']
				
	set_column_width(table.columns[0], Cm(9.5))		
	document.add_paragraph()
	
	document.add_paragraph('The following section summarizes the configuration done for Cookie protection.')
	table = document.add_table(rows=1, cols=8)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Check Signatures'
	table.cell(0,2).text = 'Signature Overides'
	table.cell(0,3).text = 'Enforcement Type'
	table.cell(0,4).text = 'Secure'
	table.cell(0,5).text = 'HTTPOnly'
	table.cell(0,6).text = 'HTTP Same Side'
	table.cell(0,7).text = 'Staging'

	
	for key in cookies:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['attackSignaturesCheck']
		cells[2].text = '\n'.join(key['signatureOverrides'])
		cells[3].text = key['enforcementType']
		cells[4].text = key['securedOverHttpsConnection']
		cells[5].text = key['accessibleOnlyThroughTheHttpProtocol']
		cells[6].text = key['insertSameSiteAttribute']
		cells[7].text = key['performStaging']

	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Cookies":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Cookies":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Cookies":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')	
	
	document.add_paragraph()
	document.add_paragraph()
			
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_redirection (document, blocking_settings, domains, suggestions):

	violations = ["Illegal redirection attempt"]
	num_of_suggestions = 0	
	
	document.add_heading('Redirection', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for Redirection Attempts violations.')
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']
					
	set_column_width(table.columns[0], Cm(9.5))		
	document.add_paragraph()
	wildcard = 0
	document.add_paragraph('The following section summarizes the configuration done for HTTP Redirection protection.')
	table = document.add_table(rows=1, cols=2)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Domain Name'
	table.cell(0,1).text = 'Include SubDomains'
	
	for key in domains:
		cells = table.add_row().cells
		cells[0].text = key['domainName']
		cells[1].text = key['includeSubdomains']

	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Redirection Domains":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Redirection Domains":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Redirection Domains":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')	
	
	document.add_paragraph()
	document.add_paragraph()
	
	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 ASM - Config Review.docx")

def word_file_methods (document, blocking_settings, methods, suggestions):
	
	violations = ["Illegal method"]
	num_of_suggestions = 0	

	document.add_heading('Allowed Methods', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for HTTP Methods violations.')
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']

				
	set_column_width(table.columns[0], Cm(9.5))		
	document.add_paragraph()

	document.add_paragraph('The following section summarizes the configuration for HTTP Method protection.')
	table = document.add_table(rows=1, cols=2)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Method Name'
	table.cell(0,1).text = 'Act as Method'
	for key in methods:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['actAsMethod']
	
	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Methods":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Methods":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Methods":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')	
	

	document.add_paragraph()
	document.add_paragraph()
		
	document.save("reports/F5 ASM - Config Review.docx")
	
def word_file_whitelist (document, blocking_settings, whitelist, suggestions):
	num_of_suggestions = 0
	violations = ["IP is blacklisted"]
	document.add_heading('IP Address List', level=3)
	document.add_paragraph('The following section summarizes the IP Address list.')

	if len(whitelist)>0:
		table = document.add_table(rows=1, cols=7)
		table.style = 'Table_rows'
		table.cell(0,0).text = 'IP Address'
		table.cell(0,1).text = 'Mask'
		table.cell(0,2).text = 'Trusted IP'
		table.cell(0,3).text = 'Never Anomalies'
		table.cell(0,4).text = 'Never Requests'
		table.cell(0,5).text = 'Ignore Anomalies'
		table.cell(0,6).text = 'Ignore Reputation'
		
		for key in whitelist:
			cells = table.add_row().cells
			cells[0].text = (key['ipAddress'])
			cells[1].text = (key['ipMask'])
			cells[2].text = (key['trustedByPolicyBuilder'])
			cells[3].text = (key['ignoreAnomalies'])
			cells[4].text = (key['neverLogRequests'])
			cells[5].text = (key['ignoreAnomalies'])
			cells[6].text = (key['ignoreIpReputation'])
			
	else:
		document.add_paragraph()
		p = document.add_paragraph()
		runner = p.add_run('The IP Address list is not configured.')
		runner.bold = True

	
	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="IP Exceptions":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="IP Exceptions":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="IP Exceptions":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')	
	
	document.add_paragraph()
	document.add_paragraph()
		
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_response (document, blocking_settings, allowed_responses, suggestions):

	violations = ["Illegal HTTP status in response"]
	num_of_suggestions = 0
	document.add_heading('HTTP Response Codes', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for HTTP Response code violations.')
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']
				
	set_column_width(table.columns[0], Cm(9.5))		
	document.add_paragraph()

	document.add_paragraph('The following section summarizes the configuration for HTTP Response code protection.')
	table = document.add_table(rows=1, cols=1)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'HTTP Response Codes'
	for key in allowed_responses:
		cells = table.add_row().cells
		cells[0].text = str(key)

		
	document.add_paragraph()
	document.add_paragraph()

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Responses":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Responses":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Responses":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')	

		
	document.add_paragraph()
	document.add_paragraph()
		
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_geolocation (document, blocking_settings, disallowed_geolocations, suggestions):		

	violations = ["Access from disallowed Geolocation"]
	num_of_suggestions = 0
	document.add_heading('Disallowed Geolocation', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for Disallowed Geolocations.')
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']
				
	set_column_width(table.columns[0], Cm(9.5))		
	document.add_paragraph()

	document.add_paragraph('The following section summarizes the configuration for HTTP Response code protection.')


	if len(disallowed_geolocations)>0:
		table = document.add_table(rows=1, cols=1)
		table.style = 'Table_rows'
		table.cell(0,0).text = 'Disallowed Countries'
		for key in disallowed_geolocations:
			cells = table.add_row().cells
			cells[0].text = key['countryName']
	else:
		document.add_paragraph()
		p = document.add_paragraph()
		runner = p.add_run('No Countries have been disallowed.')
		runner.bold = True
		
	document.add_paragraph()
	document.add_paragraph()
	
	for key in suggestions:
		if key['severity'] == "error" and key['section']=="Geolocation":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="Geolocation":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="Geolocation":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')	

		
	document.add_paragraph()
	document.add_paragraph()
		
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_ipi (document, blocking_settings, ipi, ipi_categories, suggestions):
	violations = ["Access from malicious IP address"]
	num_of_suggestions = 0
	document.add_heading('IP Intelligence', level=3)
	document.add_paragraph('The following section shows the blocking settings configuration for IP Intelligence.')
	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_column'
	table.cell(0,0).text = 'Violation Name'
	table.cell(0,1).text = 'Learn'
	table.cell(0,2).text = 'Alarm'
	table.cell(0,3).text = 'Block'
	
	for key in blocking_settings:
		if key['name'] in violations:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['learn']
			cells[2].text = key['alarm']
			cells[3].text = key['block']
				
	set_column_width(table.columns[0], Cm(9.5))		
	document.add_paragraph()

	document.add_paragraph('The following section summarizes the configuration for HTTP Response code protection.')
	p = document.add_paragraph()
	runner = p.add_run("IP Intelligence Status:	")
	runner.bold = True
	if ipi=='"Yes"':
		p.add_run ("Enabled")
	else:
		p.add_run ("Disabled")
		
	document.add_paragraph()	
	
	if len(ipi_categories)>0:
		table = document.add_table(rows=1, cols=3)
		table.style = 'Table_rows'
		table.cell(0,0).text = 'IP Intelligence Categories'
		for key in ipi_categories:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['alarm']
			cells[2].text = key['block']	
	else:
		document.add_paragraph()
		p = document.add_paragraph()
		runner = p.add_run('The IP Intelligence Categories are not being enabled.')
		runner.bold = True
	
	document.add_paragraph()
	document.add_paragraph()
	

	for key in suggestions:
		if key['severity'] == "error" and key['section']=="IP Intelligence":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "warning" and key['section']=="IP Intelligence":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['severity'] == "info" and key['section']=="IP Intelligence":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
	for key in suggestions:
		if key['txt'] in violations and key['section']=="Blocking Settings":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
				
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')	

		
	document.add_paragraph()
	document.add_paragraph()
		
	document.save("reports/F5 ASM - Config Review.docx")

def word_file_summary (document, suggestions,names):
	num_of_suggestions = 0
	table = document.add_table(rows=1, cols=4)

	table.style = 'Table Grid'
	table.cell(0,0).text = '#'
	table.cell(0,1).text = 'Severity'
	table.cell(0,2).text = 'Suggestion'
	table.cell(0,3).text = 'Category'
	
	for key in suggestions:
		if key['severity']=="error":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']
			
	for key in suggestions:
		if key['severity']=="warning":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']
			
	for key in suggestions:
		if key['severity']=="info":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']			
				
	set_column_width(table.columns[0], Cm(1))		
	set_column_width(table.columns[1], Cm(1))		
	set_column_width(table.columns[2], Cm(12))		
	set_column_width(table.columns[3], Cm(3))		
	document.add_paragraph()
		
	document.add_paragraph()
	document.add_paragraph()
		
	document.save("reports/"+names+"_summary.docx")
	